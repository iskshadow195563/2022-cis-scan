import sys
import os
import json
import datetime
import warnings
from PyQt5.QtWidgets import (QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                             QPushButton, QLabel, QTableWidget, QTableWidgetItem,
                             QComboBox, QHeaderView, QFileDialog, QMessageBox, QDialog, QDialogButtonBox, QApplication, QMenu)
from PyQt5.QtGui import QColor, QPainter, QFont, QIcon
from PyQt5.QtPrintSupport import QPrintDialog, QPrinter
from PyQt5.QtCore import Qt, QTimer, QSettings
import matplotlib
import matplotlib.font_manager
from matplotlib.figure import Figure
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas

_CJK_FONT_CANDIDATES = [
    "Microsoft YaHei",
    "SimHei",
    "Microsoft JhengHei",
    "WenQuanYi Zen Hei",
    "Noto Sans CJK SC",
    "Noto Sans CJK HK",
]
_available_fonts = [f.name for f in matplotlib.font_manager.fontManager.ttflist]
_cjk_font = next((f for f in _CJK_FONT_CANDIDATES if f in _available_fonts), None)
if _cjk_font:
    matplotlib.rcParams["font.sans-serif"] = [_cjk_font] + matplotlib.rcParams["font.sans-serif"]
matplotlib.rcParams["axes.unicode_minus"] = False
warnings.filterwarnings("ignore", message=r"Glyph .* missing from font", category=UserWarning)
from core.language_manager import tr, lang_manager
from core.report_compare import build_comparison_rows
from core.report_status import (
    FAIL_STATUS,
    PASS_STATUS,
    ERROR_STATUS,
    UNSUPPORTED_STATUS,
    MISSING_SCRIPT_STATUS,
    SCRIPT_ERROR_STATUS,
    normalize_report_data,
    normalize_report_status,
)
from core.scanner import build_suggestion
from gui.background_manager import BackgroundWidget, BackgroundManager
from gui.theme_tokens import build_theme_stylesheet, resolve_theme_name

class ReportWindow(QMainWindow):
    def __init__(self, report_path):
        super().__init__()
        self.report_path = report_path
        try:
            self.report_data = self.load_report(report_path)
            self.load_styles()
            self.init_ui()
            self.setObjectName("ReportWindow")
            self.restore_window_geometry()
        except Exception as e:
            print(f"Error initializing ReportWindow: {e}")
            # We should ideally show a message box here, but since this is a window init,
            # it's better to catch it in the caller.

    def load_report(self, path):
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        normalize_report_data(data)
        self.ensure_suggestions(data)
        return data

    def load_styles(self):
        theme_name = resolve_theme_name(QApplication.instance())
        self.setStyleSheet(build_theme_stylesheet(theme_name))

    def init_ui(self):
        self.setWindowTitle(tr("report_title"))
        self.setMinimumSize(1000, 800)

        self._bg_manager = BackgroundManager()
        central_widget = BackgroundWidget(self._bg_manager)
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)

        # Header: Date/Time, Scan Mode, Server IP and Score
        header_layout = QHBoxLayout()
        scan_info = self.report_data.get('scan_info', {})
        date_str = scan_info.get('date', 'N/A')
        time_str = scan_info.get('time', 'N/A')

        date_disp = lang_manager.format_date(date_str)
        time_disp = lang_manager.format_time(time_str)
        info_label = QLabel(f"{tr('date')}: {date_disp}  {tr('time')}: {time_disp}")
        info_label.setStyleSheet("font-size: 15px;")
        header_left = QWidget()
        header_left_layout = QVBoxLayout(header_left)
        header_left_layout.setContentsMargins(0, 0, 0, 0)
        header_left_layout.addWidget(info_label)

        # Scan mode display
        raw_mode = scan_info.get('scan_mode', 'local')
        mode_display = tr(f"scan_mode_{raw_mode}")
        self.scan_mode_label = QLabel(f"{tr('report_scan_mode')}: {mode_display}")
        self.scan_mode_label.setStyleSheet("font-size: 15px;")
        header_left_layout.addWidget(self.scan_mode_label)

        # Server IP display
        raw_target = scan_info.get('remote_target', '')
        if raw_mode == "remote" and raw_target:
            server_ip = raw_target
        else:
            server_ip = scan_info.get('local_ip', '')
        self.server_ip_label = QLabel(f"{tr('report_server_ip')}: {server_ip}")
        self.server_ip_label.setStyleSheet("font-size: 15px;")
        header_left_layout.addWidget(self.server_ip_label)

        self.nowtime_label = QLabel("")
        self.nowtime_label.setStyleSheet("font-size: 15px;")
        header_left_layout.addWidget(self.nowtime_label)
        header_layout.addWidget(header_left)

        header_layout.addStretch()

        score = self.calculate_score()
        self.score_label = QLabel(f"{tr('score')}: {score:.2f}%")
        self.score_label.setStyleSheet("font-size: 24px; font-weight: bold; color: #2c3e50;")
        header_layout.addWidget(self.score_label)

        main_layout.addLayout(header_layout)

        # Inline alert banner
        self.alert_bar = QWidget()
        self.alert_bar.setProperty("card", True)
        self.alert_bar.setVisible(False)
        _alert_layout = QHBoxLayout(self.alert_bar)
        _alert_layout.setContentsMargins(10, 6, 10, 6)
        self.alert_label = QLabel("")
        self.alert_label.setWordWrap(True)
        self.alert_close = QPushButton("×")
        self.alert_close.setFixedWidth(28)
        self.alert_close.clicked.connect(lambda: self.alert_bar.setVisible(False))
        _alert_layout.addWidget(self.alert_label, 1)
        _alert_layout.addWidget(self.alert_close, 0)
        main_layout.addWidget(self.alert_bar)

        self.update_nowtime()
        self._nowtime_timer = QTimer(self)
        self._nowtime_timer.timeout.connect(self.update_nowtime)
        self._nowtime_timer.start(1000)

        # Middle: Pie Chart and Summary
        middle_layout = QHBoxLayout()

        # Pie Chart - Use OO API
        self.figure = Figure(figsize=(5, 5))
        self.canvas = FigureCanvas(self.figure)
        self.ax = self.figure.add_subplot(111)
        self.update_pie_chart()
        middle_layout.addWidget(self.canvas, 1)

        # Summary and Comparison Controls
        right_panel = QVBoxLayout()

        summary_layout = QVBoxLayout()
        summary_title = QLabel(tr("summary"))
        summary_title.setStyleSheet("font-size: 16px; font-weight: bold;")
        summary_layout.addWidget(summary_title)
        self.summary_labels = []
        for label_text, count_value, bg_color in self.build_summary_rows():
            label = QLabel(f"{label_text}: {count_value}")
            label.setStyleSheet(f"padding: 6px 10px; border-radius: 6px; background-color: {bg_color};")
            self.summary_labels.append(label)
            summary_layout.addWidget(label)
        right_panel.addLayout(summary_layout)

        # Comparison
        comp_layout = QVBoxLayout()
        comp_layout.addWidget(QLabel(tr("compare_previous")))
        self.comp_combo = QComboBox()
        self.populate_comparison_files()
        comp_layout.addWidget(self.comp_combo)

        self.btn_compare = QPushButton(tr("compare"))
        self.btn_compare.clicked.connect(self.compare_reports)
        comp_layout.addWidget(self.btn_compare)
        right_panel.addLayout(comp_layout)

        right_panel.addStretch()

        export_layout = QVBoxLayout()
        export_title = QLabel(tr("export_report"))
        export_title.setStyleSheet("font-size: 16px; font-weight: bold;")
        export_layout.addWidget(export_title)

        self.btn_export = QPushButton(tr("export_report"))
        self.btn_export.setMinimumHeight(32)
        export_menu = QMenu(self)
        export_menu.addAction(tr("export_docx"), self.export_docx)
        export_menu.addAction(tr("export_txt"), self.export_txt)
        export_menu.addAction(tr("export_xlsx"), self.export_xlsx)
        export_menu.addAction(tr("export_csv"), self.export_csv)
        self.btn_export.setMenu(export_menu)
        export_layout.addWidget(self.btn_export)

        right_panel.addLayout(export_layout)

        # Unsupported/Error details button
        self.unsupported_layout = QVBoxLayout()
        unsupported_title = QLabel(tr("filter_unsupported"))
        unsupported_title.setStyleSheet("font-size: 16px; font-weight: bold;")
        self.unsupported_layout.addWidget(unsupported_title)
        self.btn_unsupported = QPushButton(tr("filter_unsupported"))
        self.btn_unsupported.clicked.connect(self.show_unsupported_details)
        self.unsupported_layout.addWidget(self.btn_unsupported)
        right_panel.addLayout(self.unsupported_layout)

        self.btn_back = QPushButton(tr("back"))
        self.btn_back.clicked.connect(self.close)
        right_panel.addWidget(self.btn_back)

        self.btn_print = QPushButton(tr("print_report"))
        self.btn_print.clicked.connect(self.print_page)
        right_panel.addWidget(self.btn_print)

        middle_layout.addLayout(right_panel, 1)
        main_layout.addLayout(middle_layout, 2)

        # Bottom: Results Table
        self.table = QTableWidget()
        self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels([tr("table_check_no"), tr("table_code"), tr("table_level"), tr("table_description"), tr("table_suggestion"), tr("table_status")])
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(2, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(3, QHeaderView.Stretch)
        header.setSectionResizeMode(4, QHeaderView.Stretch)
        header.setSectionResizeMode(5, QHeaderView.ResizeToContents)
        self.table.verticalHeader().setVisible(False)
        self.table.setAlternatingRowColors(True)
        self.table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.populate_table()
        main_layout.addWidget(self.table, 3)

    def show_alert(self, message: str, is_error: bool = False):
        color = "#e74c3c" if is_error else "#2ecc71"
        bg = "#fdecea" if is_error else "#ecf9f1"
        self.alert_label.setText(message)
        self.alert_bar.setStyleSheet(
            f"background-color: {bg}; border: 1px solid {color}; border-radius: 6px; color: #2c3e50;"
        )
        self.alert_bar.setVisible(True)

    def update_nowtime(self):
        now = datetime.datetime.now()
        date_disp = lang_manager.format_date(now.strftime("%Y-%m-%d"))
        time_disp = lang_manager.format_time(now.strftime("%H:%M:%S"))
        self.nowtime_label.setText(f"{tr('nowtime')}: {date_disp} {time_disp}")

    def get_stats(self):
        results = self.report_data.get("results", [])
        stats = {
            PASS_STATUS: 0,
            FAIL_STATUS: 0,
            MISSING_SCRIPT_STATUS: 0,
            ERROR_STATUS: 0,
            SCRIPT_ERROR_STATUS: 0,
            UNSUPPORTED_STATUS: 0,
        }
        for item in results:
            status = normalize_report_status(item.get("status"))
            if status in stats:
                stats[status] += 1
        stats["Total"] = len(results)
        return stats

    def translate_status(self, status):
        normalized = normalize_report_status(status)
        mapping = {
            PASS_STATUS: tr("status_pass"),
            FAIL_STATUS: tr("status_fail"),
            ERROR_STATUS: tr("status_error"),
            UNSUPPORTED_STATUS: tr("status_not_supported"),
            MISSING_SCRIPT_STATUS: tr("status_not_checked"),
            SCRIPT_ERROR_STATUS: tr("status_error"),
        }
        return mapping.get(normalized, str(normalized))

    def status_color(self, status):
        colors = {
            PASS_STATUS: "#d4edda",
            FAIL_STATUS: "#f8d7da",
            ERROR_STATUS: "#fff3cd",
            UNSUPPORTED_STATUS: "#d6d8db",
            MISSING_SCRIPT_STATUS: "#e2e3e5",
            SCRIPT_ERROR_STATUS: "#fdecea",
        }
        return colors.get(normalize_report_status(status), "#e2e3e5")

    def translate_compare_status(self, status):
        if status is None:
            return tr("compare_status_missing")
        return self.translate_status(status)

    def comparison_status_color(self, status):
        if status is None:
            return "#f1f3f5"
        return self.status_color(status)

    def comparison_row_changed(self, row):
        if not isinstance(row, dict):
            return False
        return bool(row.get("changed") or row.get("old_status") != row.get("new_status"))

    def format_comparison_name(self, row):
        if not isinstance(row, dict):
            return ""
        code = str(row.get("code") or "").strip()
        name = str(row.get("name") or "").strip()
        if code and name and code not in name:
            return f"{code} - {name}"
        return name or code

    def format_comparison_detail(self, row, prefix):
        if not isinstance(row, dict):
            return "-"
        parts = []
        detail = str(row.get(f"{prefix}_detail") or "").strip()
        expected = str(row.get(f"{prefix}_expected_value") or "").strip()
        actual = str(row.get(f"{prefix}_actual_value") or "").strip()
        if detail:
            parts.append(detail)
        if expected:
            parts.append(f"{tr('csv_expected_value')}: {expected}")
        if actual:
            parts.append(f"{tr('csv_actual_value')}: {actual}")
        return "\n".join(parts) if parts else "-"

    def build_summary_rows(self):
        stats = self.get_stats()
        return [
            (tr("total_items"), stats["Total"], "#ecf0f1"),
            (tr("status_pass"), stats[PASS_STATUS], self.status_color(PASS_STATUS)),
            (tr("status_fail"), stats[FAIL_STATUS], self.status_color(FAIL_STATUS)),
            (tr("status_not_supported"), stats[UNSUPPORTED_STATUS], self.status_color(UNSUPPORTED_STATUS)),
            (tr("status_not_checked"), stats[MISSING_SCRIPT_STATUS], self.status_color(MISSING_SCRIPT_STATUS)),
            (tr("status_error"), stats[ERROR_STATUS] + stats.get(SCRIPT_ERROR_STATUS, 0), self.status_color(ERROR_STATUS)),
        ]

    def get_pie_counts(self):
        stats = self.get_stats()
        pass_count = stats.get(PASS_STATUS, 0)
        fail_count = sum(
            stats.get(status, 0)
            for status in (
                FAIL_STATUS,
                UNSUPPORTED_STATUS,
                MISSING_SCRIPT_STATUS,
                ERROR_STATUS,
                SCRIPT_ERROR_STATUS,
            )
        )
        return pass_count, fail_count

    def calculate_score(self):
        pass_count, fail_count = self.get_pie_counts()
        total = pass_count + fail_count
        if total == 0:
            return 0
        return (pass_count / total) * 100

    def update_pie_chart(self):
        pass_count, fail_count = self.get_pie_counts()
        pie_rows = [
            (PASS_STATUS, pass_count),
            (FAIL_STATUS, fail_count),
        ]
        labels = [
            f"{self.translate_status(status)} ({value})"
            for status, value in pie_rows
            if value > 0
        ]
        sizes = [value for _status, value in pie_rows if value > 0]
        colors = [self.status_color(status) for status, value in pie_rows if value > 0]
        self.ax.clear()
        if sizes:
            self.ax.pie(
                sizes,
                labels=None,
                autopct='%1.1f%%',
                startangle=140,
                colors=colors,
                wedgeprops={"linewidth": 1, "edgecolor": "white"}
            )
            self.ax.legend(labels, loc="center left", bbox_to_anchor=(1.0, 0.5), frameon=False)
        else:
            self.ax.text(0, 0.12, tr("no_report_data"), ha="center", va="center", fontsize=12, color="#7f8c8d")
        score = self.calculate_score()
        self.ax.text(0, -0.08 if not sizes else 0, f"{score:.1f}%", ha="center", va="center", fontsize=16, color="#2c3e50")
        self.ax.axis('equal')
        self.figure.tight_layout()
        with warnings.catch_warnings():
            warnings.filterwarnings("ignore", message=r"Glyph .* missing from font", category=UserWarning)
            self.canvas.draw()

    def populate_table(self):
        results = self.report_data["results"]
        self.table.setRowCount(len(results))
        for i, res in enumerate(results):
            check_item = QTableWidgetItem(str(i + 1))
            check_item.setTextAlignment(Qt.AlignCenter)
            self.table.setItem(i, 0, check_item)

            self.table.setItem(i, 1, QTableWidgetItem(res.get("code", "")))
            level_item = QTableWidgetItem(res.get("level", ""))
            level_item.setTextAlignment(Qt.AlignCenter)
            self.table.setItem(i, 2, level_item)
            self.table.setItem(i, 3, QTableWidgetItem(res.get("description", "")))
            self.table.setItem(i, 4, QTableWidgetItem(self.get_suggestion(res)))
            status_raw = normalize_report_status(res.get("status"))
            status_item = QTableWidgetItem(self.translate_status(status_raw))
            status_item.setData(Qt.UserRole, status_raw)
            status_item.setBackground(QColor(self.status_color(status_raw)))
            status_item.setTextAlignment(Qt.AlignCenter)
            self.table.setItem(i, 5, status_item)

    def show_unsupported_details(self):
        results = self.report_data.get("results", [])
        unsupported_items = []
        error_items = []
        for res in results:
            status_raw = normalize_report_status(res.get("status"))
            if status_raw == UNSUPPORTED_STATUS:
                unsupported_items.append(res)
            elif status_raw in (ERROR_STATUS, SCRIPT_ERROR_STATUS):
                error_items.append(res)

        dialog = QDialog(self)
        dialog.setWindowTitle(tr("unsupported_error_title"))
        dialog.setMinimumSize(750, 450)
        layout = QVBoxLayout(dialog)

        count_label = QLabel(tr("unsupported_error_count", len(unsupported_items), len(error_items)))
        count_label.setStyleSheet("font-size: 14px; font-weight: bold; padding: 6px;")
        layout.addWidget(count_label)

        table = QTableWidget()
        table.setColumnCount(4)
        table.setHorizontalHeaderLabels([tr("table_code"), tr("table_level"), tr("table_description"), tr("table_status")])
        header = table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(2, QHeaderView.Stretch)
        header.setSectionResizeMode(3, QHeaderView.ResizeToContents)
        table.verticalHeader().setVisible(False)
        table.setAlternatingRowColors(True)
        table.setEditTriggers(QTableWidget.NoEditTriggers)
        table.setSelectionBehavior(QTableWidget.SelectRows)

        items = unsupported_items + error_items
        table.setRowCount(len(items))
        for i, res in enumerate(items):
            table.setItem(i, 0, QTableWidgetItem(res.get("code", "")))
            level_item = QTableWidgetItem(res.get("level", ""))
            level_item.setTextAlignment(Qt.AlignCenter)
            table.setItem(i, 1, level_item)

            detail = res.get("detail", "")
            desc = res.get("description", "")
            display_text = desc
            if detail:
                display_text = f"{desc}  [{detail}]"
            table.setItem(i, 2, QTableWidgetItem(display_text))

            status_raw = normalize_report_status(res.get("status"))
            status_item = QTableWidgetItem(self.translate_status(status_raw))
            status_item.setData(Qt.UserRole, status_raw)
            status_item.setBackground(QColor(self.status_color(status_raw)))
            status_item.setTextAlignment(Qt.AlignCenter)
            table.setItem(i, 3, status_item)

        layout.addWidget(table, 1)

        buttons = QDialogButtonBox(QDialogButtonBox.Close)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)

        dialog.exec_()

    def ensure_suggestions(self, report_data):
        results = (report_data or {}).get("results") or []
        for res in results:
            if not isinstance(res, dict):
                continue
            if res.get("suggestion"):
                continue
            res["suggestion"] = build_suggestion(res.get("code", ""), res.get("description", ""))

    def get_suggestion(self, res):
        if not isinstance(res, dict):
            return ""
        suggestion = res.get("suggestion")
        if suggestion:
            return suggestion
        return build_suggestion(res.get("code", ""), res.get("description", ""))

    def restore_window_geometry(self):
        settings = QSettings("project-001", "WindowsSecurityAuditor")
        geometry = settings.value("report_window/geometry")
        if geometry:
            try:
                self.restoreGeometry(geometry)
            except Exception:
                pass

    def closeEvent(self, event):
        try:
            settings = QSettings("project-001", "WindowsSecurityAuditor")
            settings.setValue("report_window/geometry", self.saveGeometry())
        except Exception:
            pass
        super().closeEvent(event)

    def populate_comparison_files(self):
        output_dir = os.path.dirname(self.report_path)
        files = [
            f for f in os.listdir(output_dir)
            if self.is_scan_report_file(os.path.join(output_dir, f)) and os.path.join(output_dir, f) != self.report_path
        ]
        files.sort(reverse=True) # Newer first
        self.comp_combo.addItems(files)

    def is_scan_report_file(self, path):
        filename = os.path.basename(path).lower()
        if not (filename.startswith("report_") and filename.endswith(".json")):
            return False
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return isinstance(data, dict) and isinstance(data.get("results"), list)
        except Exception:
            return False

    def compare_reports(self):
        selected_file = self.comp_combo.currentText()
        if not selected_file:
            return

        prev_path = os.path.join(os.path.dirname(self.report_path), selected_file)
        try:
            with open(prev_path, "r", encoding="utf-8") as f:
                prev_data = json.load(f)
        except Exception as e:
            QMessageBox.critical(self, tr("compare"), str(e))
            return

        rows = build_comparison_rows(prev_data, self.report_data)
        self.show_comparison_results(rows, selected_file)

    def show_comparison_results(self, rows, selected_filename):
        dialog = QDialog(self)
        dialog.setWindowTitle(tr("compare_result_title", selected_filename))
        dialog.setMinimumSize(900, 520)

        layout = QVBoxLayout(dialog)

        filter_layout = QHBoxLayout()
        filter_layout.addWidget(QLabel(tr("compare_filter_label")))
        self._compare_filter = "both"

        btn_both = QPushButton(tr("compare_filter_both"))
        btn_both.clicked.connect(lambda: self._apply_compare_filter(table, safe_rows, "both", [btn_both, btn_matched, btn_mismatched]))
        btn_matched = QPushButton(tr("compare_filter_matched"))
        btn_matched.clicked.connect(lambda: self._apply_compare_filter(table, safe_rows, "matched", [btn_both, btn_matched, btn_mismatched]))
        btn_mismatched = QPushButton(tr("compare_filter_mismatched"))
        btn_mismatched.clicked.connect(lambda: self._apply_compare_filter(table, safe_rows, "mismatched", [btn_both, btn_matched, btn_mismatched]))

        for b in [btn_both, btn_matched, btn_mismatched]:
            b.setCheckable(True)
        btn_both.setChecked(True)

        filter_layout.addWidget(btn_both)
        filter_layout.addWidget(btn_matched)
        filter_layout.addWidget(btn_mismatched)
        filter_layout.addStretch()
        layout.addLayout(filter_layout)

        table = QTableWidget()
        table.setColumnCount(5)
        table.setHorizontalHeaderLabels([
            tr("compare_col_name"),
            tr("compare_col_old"),
            tr("compare_col_old_detail"),
            tr("compare_col_new"),
            tr("compare_col_new_detail"),
        ])
        table.verticalHeader().setVisible(False)
        table.setAlternatingRowColors(True)
        table.setEditTriggers(QTableWidget.NoEditTriggers)
        table.setSelectionBehavior(QTableWidget.SelectRows)
        table.setWordWrap(True)

        header = table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(2, QHeaderView.Stretch)
        header.setSectionResizeMode(3, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(4, QHeaderView.Stretch)
        header.setMinimumSectionSize(90)

        safe_rows = rows or []
        table.setRowCount(len(safe_rows))
        bold = QFont()
        bold.setBold(True)

        for i, r in enumerate(safe_rows):
            name = self.format_comparison_name(r)
            old_raw = r.get("old_status") if isinstance(r, dict) else None
            new_raw = r.get("new_status") if isinstance(r, dict) else None
            changed = self.comparison_row_changed(r)

            name_item = QTableWidgetItem(str(name))
            if changed:
                name_item.setFont(bold)
                name_item.setBackground(QColor("#fffacd"))
            table.setItem(i, 0, name_item)

            old_item = QTableWidgetItem(self.translate_compare_status(old_raw))
            old_item.setData(Qt.UserRole, old_raw)
            old_item.setTextAlignment(Qt.AlignCenter)
            old_item.setBackground(QColor(self.comparison_status_color(old_raw)))
            if changed:
                old_item.setForeground(QColor("#2c3e50"))
            table.setItem(i, 1, old_item)

            old_detail_item = QTableWidgetItem(self.format_comparison_detail(r, "old"))
            old_detail_item.setToolTip(old_detail_item.text())
            if changed:
                old_detail_item.setBackground(QColor("#fffdf2"))
            table.setItem(i, 2, old_detail_item)

            new_item = QTableWidgetItem(self.translate_compare_status(new_raw))
            new_item.setData(Qt.UserRole, new_raw)
            new_item.setTextAlignment(Qt.AlignCenter)
            new_item.setBackground(QColor(self.comparison_status_color(new_raw)))
            if changed:
                new_item.setForeground(QColor("blue"))
            table.setItem(i, 3, new_item)

            new_detail_item = QTableWidgetItem(self.format_comparison_detail(r, "new"))
            new_detail_item.setToolTip(new_detail_item.text())
            if changed:
                new_detail_item.setBackground(QColor("#fffdf2"))
            table.setItem(i, 4, new_detail_item)

        table.resizeRowsToContents()

        layout.addWidget(table, 1)

        buttons = QDialogButtonBox(QDialogButtonBox.Close)
        buttons.rejected.connect(dialog.reject)
        buttons.accepted.connect(dialog.accept)
        layout.addWidget(buttons)

        dialog.exec_()

    def _apply_compare_filter(self, table, rows, filter_mode, buttons):
        for b in buttons:
            b.setChecked(b is buttons[{"both": 0, "matched": 1, "mismatched": 2}[filter_mode]])
        for row_idx, r in enumerate(rows):
            changed = self.comparison_row_changed(r)
            if filter_mode == "both":
                table.setRowHidden(row_idx, False)
            elif filter_mode == "matched":
                table.setRowHidden(row_idx, changed)
            elif filter_mode == "mismatched":
                table.setRowHidden(row_idx, not changed)

    def print_page(self):
        printer = QPrinter(QPrinter.HighResolution)
        dialog = QPrintDialog(printer, self)
        if dialog.exec_() == QPrintDialog.Accepted:
            self.render(printer)

    def default_export_path(self, extension):
        base_name = os.path.splitext(os.path.basename(self.report_path))[0]
        filename = f"{base_name}.{extension}"
        return os.path.join(os.path.dirname(self.report_path), filename)

    def export_docx(self):
        default_path = self.default_export_path("docx")
        filepath, _ = QFileDialog.getSaveFileName(self, tr("export_report"), default_path, "Word Document (*.docx)")
        if not filepath:
            return
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(tr("report_title"), level=1)
            scan_info = self.report_data.get("scan_info", {})
            date_str = scan_info.get("date", "N/A")
            time_str = scan_info.get("time", "N/A")
            score = self.calculate_score()
            doc.add_paragraph(f"{tr('date')}: {date_str}")
            doc.add_paragraph(f"{tr('time')}: {time_str}")
            doc.add_paragraph(f"{tr('score')}: {score:.2f}%")

            doc.add_paragraph("")
            doc.add_heading(tr("summary"), level=2)
            summary_table = doc.add_table(rows=1, cols=2)
            header_cells = summary_table.rows[0].cells
            header_cells[0].text = tr("item")
            header_cells[1].text = tr("count")
            for label_text, count_value, _ in self.build_summary_rows():
                row_cells = summary_table.add_row().cells
                row_cells[0].text = label_text
                row_cells[1].text = str(count_value)

            doc.add_paragraph("")
            doc.add_heading(tr("report_title"), level=2)
            table = doc.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = tr("table_check_no")
            hdr_cells[1].text = tr("table_code")
            hdr_cells[2].text = tr("table_level")
            hdr_cells[3].text = tr("table_description")
            hdr_cells[4].text = tr("table_suggestion")
            hdr_cells[5].text = tr("table_status")
            for idx, res in enumerate(self.report_data.get("results", []), start=1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = res.get("code", "")
                row_cells[2].text = res.get("level", "")
                row_cells[3].text = res.get("description", "")
                row_cells[4].text = self.get_suggestion(res)
                row_cells[5].text = self.translate_status(res.get("status"))
            doc.save(filepath)
            self.show_alert(tr("export_success"))
        except Exception as e:
            self.show_alert(f"{tr('export_failed')}: {e}", is_error=True)


    def export_txt(self):
        default_path = self.default_export_path("txt")
        filepath, _ = QFileDialog.getSaveFileName(self, tr("export_report"), default_path, "Text File (*.txt)")
        if not filepath:
            return
        try:
            scan_info = self.report_data.get("scan_info", {})
            date_str = scan_info.get("date", "N/A")
            time_str = scan_info.get("time", "N/A")
            score = self.calculate_score()
            lines = []
            lines.append(tr("report_title"))
            lines.append(f"{tr('date')}: {date_str}    {tr('time')}: {time_str}")
            lines.append(f"{tr('score')}: {score:.2f}%")
            lines.append("")
            lines.append(tr("summary"))
            for label_text, count_value, _ in self.build_summary_rows():
                lines.append(f"{label_text}: {count_value}")
            lines.append("")
            headers = [tr("table_check_no"), tr("table_code"), tr("table_level"), tr("table_description"), tr("table_suggestion"), tr("table_status")]
            no_w, code_w, level_w, desc_w, suggestion_w, status_w = 6, 12, 8, 48, 48, 16
            header_line = f"{headers[0]:<{no_w}} {headers[1]:<{code_w}} {headers[2]:<{level_w}} {headers[3]:<{desc_w}} {headers[4]:<{suggestion_w}} {headers[5]:<{status_w}}"
            lines.append(header_line)
            lines.append("-" * len(header_line))
            import textwrap
            for row_no, res in enumerate(self.report_data.get("results", []), start=1):
                code = res.get("code", "")
                level = res.get("level", "")
                desc = res.get("description", "")
                suggestion = self.get_suggestion(res)
                status = self.translate_status(res.get("status"))
                wrapped_desc = textwrap.wrap(desc, width=desc_w) or [""]
                wrapped_suggestion = textwrap.wrap(suggestion, width=suggestion_w) or [""]
                total_lines = max(len(wrapped_desc), len(wrapped_suggestion))
                for line_idx in range(total_lines):
                    desc_line = wrapped_desc[line_idx] if line_idx < len(wrapped_desc) else ""
                    suggestion_line = wrapped_suggestion[line_idx] if line_idx < len(wrapped_suggestion) else ""
                    if line_idx == 0:
                        lines.append(f"{str(row_no):<{no_w}} {code:<{code_w}} {level:<{level_w}} {desc_line:<{desc_w}} {suggestion_line:<{suggestion_w}} {status:<{status_w}}")
                    else:
                        lines.append(f"{'':<{no_w}} {'':<{code_w}} {'':<{level_w}} {desc_line:<{desc_w}} {suggestion_line:<{suggestion_w}} {'':<{status_w}}")
            with open(filepath, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            self.show_alert(tr("export_success"))
        except Exception as e:
            self.show_alert(f"{tr('export_failed')}: {e}", is_error=True)

    def export_xlsx(self):
        default_path = self.default_export_path("xlsx")
        filepath, _ = QFileDialog.getSaveFileName(self, tr("export_report"), default_path, "Excel Workbook (*.xlsx)")
        if not filepath:
            return
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment
            wb = Workbook()
            ws_summary = wb.active
            ws_summary.title = tr("summary")
            bold = Font(bold=True)
            ws_summary["A1"] = tr("report_title")
            ws_summary["A1"].font = Font(bold=True, size=14)
            scan_info = self.report_data.get("scan_info", {})
            date_str = scan_info.get("date", "N/A")
            time_str = scan_info.get("time", "N/A")
            score = self.calculate_score()
            ws_summary["A3"] = tr("date")
            ws_summary["B3"] = date_str
            ws_summary["A4"] = tr("time")
            ws_summary["B4"] = time_str
            ws_summary["A5"] = tr("score")
            ws_summary["B5"] = f"{score:.2f}%"
            ws_summary["A7"] = tr("summary")
            ws_summary["A7"].font = bold
            ws_summary["A8"] = tr("item")
            ws_summary["B8"] = tr("count")
            ws_summary["A8"].font = bold
            ws_summary["B8"].font = bold
            row = 9
            for label_text, count_value, _ in self.build_summary_rows():
                ws_summary[f"A{row}"] = label_text
                ws_summary[f"B{row}"] = count_value
                row += 1
            ws_summary.column_dimensions["A"].width = 30
            ws_summary.column_dimensions["B"].width = 15

            ws_results = wb.create_sheet(tr("report_title"))
            headers = [tr("table_check_no"), tr("table_code"), tr("table_level"), tr("table_description"), tr("table_suggestion"), tr("table_status")]
            ws_results.append(headers)
            for col in range(1, 7):
                ws_results.cell(row=1, column=col).font = bold
                ws_results.cell(row=1, column=col).alignment = Alignment(horizontal="center")
            for idx, res in enumerate(self.report_data.get("results", []), start=1):
                ws_results.append([
                    idx,
                    res.get("code", ""),
                    res.get("level", ""),
                    res.get("description", ""),
                    self.get_suggestion(res),
                    self.translate_status(res.get("status"))
                ])
            ws_results.column_dimensions["A"].width = 8
            ws_results.column_dimensions["B"].width = 14
            ws_results.column_dimensions["C"].width = 8
            ws_results.column_dimensions["D"].width = 60
            ws_results.column_dimensions["E"].width = 60
            ws_results.column_dimensions["F"].width = 18
            wb.save(filepath)
            self.show_alert(tr("export_success"))
        except Exception as e:
            self.show_alert(f"{tr('export_failed')}: {e}", is_error=True)

    def export_csv(self):
        default_path = self.default_export_path("csv")
        filepath, _ = QFileDialog.getSaveFileName(self, tr("export_csv"), default_path, "CSV File (*.csv)")
        if not filepath:
            return
        try:
            import csv
            with open(filepath, "w", encoding="utf-8-sig", newline="") as f:
                writer = csv.writer(f)
                writer.writerow([
                    tr("table_check_no"),
                    tr("table_code"),
                    tr("table_level"),
                    tr("table_description"),
                    tr("table_suggestion"),
                    tr("table_status"),
                    tr("table_detail"),
                    tr("csv_expected_value"),
                    tr("csv_actual_value"),
                ])
                for idx, res in enumerate(self.report_data.get("results", []), start=1):
                    writer.writerow([
                        idx,
                        res.get("code", ""),
                        res.get("level", ""),
                        res.get("description", ""),
                        self.get_suggestion(res),
                        self.translate_status(res.get("status")),
                        res.get("status_detail") or res.get("detail", ""),
                        res.get("expected_value", ""),
                        res.get("actual_value", ""),
                    ])
            self.show_alert(tr("export_success"))
        except Exception as e:
            self.show_alert(f"{tr('export_failed')}: {e}", is_error=True)
