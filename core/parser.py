import json
import os
import re
from docx import Document

def parse_cis_docx(file_path):
    if not os.path.exists(file_path):
        return []

    doc = Document(file_path)
    items = []

    # Regex to match patterns like "17.5.4 (L1) Ensure Audit Logon = Success and Failure"
    # Or "18.11.1.1 (L1) Ensure Prevent..."
    pattern = re.compile(r'^([\d\.]+)\s+\((L\d(?:\.\d+)?)\)\s+(.+)$')

    all_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    all_text.extend([t.strip() for t in cell.text.split('\n') if t.strip()])

    for text in all_text:
        match = pattern.match(text)
        if match:
            code = match.group(1)
            level = match.group(2)
            description = match.group(3)
            items.append({
                "code": code,
                "level": level,
                "description": description,
                "full_text": text
            })
        else:
            # Try a looser match for items that might not perfectly follow the pattern
            # e.g. "2.2.37 (L2) Log on as a batch job = ..."
            loose_match = re.search(r'([\d\.]+)\s+\((L\d(?:\.\d+)?)\)', text)
            if loose_match:
                code = loose_match.group(1)
                level = loose_match.group(2)
                # Extract the rest of the text after the level
                desc_start = text.find(level) + len(level) + 1
                description = text[desc_start:].strip()
                items.append({
                    "code": code,
                    "level": level,
                    "description": description,
                    "full_text": text
                })

    return items

if __name__ == "__main__":
    docx_path = r"E:\.school\HKIIT\4116M-FYP\TEST\mini project test-001\scan check item.docx"
    items = parse_cis_docx(docx_path)

    output_path = r"E:\.school\HKIIT\4116M-FYP\data\cis_items.json"
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(items, f, indent=2, ensure_ascii=False)

    print(f"Parsed {len(items)} items and saved to {output_path}")
    # Print levels found to check for L1.11
    levels = set(item['level'] for item in items)
    print(f"Levels found: {levels}")
